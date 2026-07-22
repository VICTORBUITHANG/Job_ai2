# job_ai2_agent/resume_reader.py

import re
from pathlib import Path

from job_ai2_agent.models import EducationItem, ResumeProfile, WorkExperience


EMAIL_RE = re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.IGNORECASE)
PHONE_RE = re.compile(r"(?:\+?1[\s.-]?)?(?:\(?\d{3}\)?[\s.-]?)\d{3}[\s.-]?\d{4}")
URL_RE = re.compile(r"https?://[^\s)]+|(?:linkedin|github)\.com/[^\s)]+", re.IGNORECASE)


def read_resume_profile(path: Path) -> ResumeProfile:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return read_docx_resume_profile(path)
    if suffix == ".pdf":
        return read_pdf_resume_profile(path)
    raise ValueError("Resume must be a PDF or DOCX file.")


def read_docx_resume_profile(path):
    if path.suffix.lower() != ".docx":
        raise ValueError("DOCX resume parser only accepts .docx files.")
    return _profile_from_text(_read_docx(path))


def read_pdf_resume_profile(path):
    if path.suffix.lower() != ".pdf":
        raise ValueError("PDF resume parser only accepts .pdf files.")
    return _profile_from_text(_read_pdf(path))


def _profile_from_text(text):
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    fields = infer_basic_profile(text)
    work_experiences = infer_work_experiences(lines)
    education_items = infer_education_items(lines)
    if work_experiences:
        _add_current_work_fields(fields, work_experiences[0])
    if education_items:
        _add_top_education_fields(fields, education_items[0])
    return ResumeProfile(
        fields=fields,
        raw_text=text,
        work_experiences=work_experiences,
        education_items=education_items,
    )


def extract_resume_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_resume_text(path)
    if suffix == ".docx":
        return extract_docx_resume_text(path)
    raise ValueError("Resume must be a PDF or DOCX file.")


def extract_docx_resume_text(path):
    if path.suffix.lower() != ".docx":
        raise ValueError("DOCX text extractor only accepts .docx files.")
    return _read_docx(path)


def extract_pdf_resume_text(path):
    if path.suffix.lower() != ".pdf":
        raise ValueError("PDF text extractor only accepts .pdf files.")
    return _read_pdf(path)


def infer_basic_profile(text: str) -> dict[str, str]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    raw_full_name = _guess_name(lines)
    first_name, last_name, preferred_name = _split_name(raw_full_name)
    full_name = _compose_legal_name(first_name, last_name)
    location = _guess_location(lines)
    city, state = _split_city_state(location)
    current_job = _guess_current_job(lines)
    top_education = _guess_top_education(lines)
    profile: dict[str, str] = {
        "legal_name": full_name,
        "full_name": full_name,
        "first_name": first_name,
        "last_name": last_name,
        "preferred_name": preferred_name,
        "email": _first_match(EMAIL_RE, text),
        "phone": _first_match(PHONE_RE, text),
        "location": location,
        "city": city,
        "state": state,
        "country": "United States of America",
        "skills": _guess_section(lines, ("skills", "technical skills", "technologies")),
        "education": _guess_section(lines, ("education",)),
        "experience": _guess_section(lines, ("experience", "work experience", "employment")),
        "summary": " ".join(lines[:8]),
    }
    profile.update(current_job)
    profile.update(top_education)
    for url in URL_RE.findall(text):
        normalized = url if url.startswith("http") else f"https://{url}"
        lowered = normalized.lower()
        if "linkedin.com" in lowered:
            profile["linkedin"] = normalized
        elif "github.com" in lowered:
            profile["github"] = normalized
        else:
            profile["website"] = normalized
    return {key: value for key, value in profile.items() if value}


def infer_work_experiences(lines: list[str]) -> list[WorkExperience]:
    section = _section_lines(
        lines,
        ("work experience", "experience", "employment", "professional experience"),
        (
            "education",
            "skills",
            "core expertise",
            "publications",
            "patents",
            "awards",
            "certifications",
            "projects",
            "selected projects",
            "selected publications",
            "selected patents",
        ),
    )
    experiences: list[WorkExperience] = []
    current: WorkExperience | None = None
    for line in section:
        if _skip_resume_noise(line):
            continue
        parsed = _parse_work_header(line)
        if parsed:
            if current:
                experiences.append(current)
            current = parsed
            continue
        if current:
            current.description = _append_description(current.description, line)
    if current:
        experiences.append(current)
    return experiences


def infer_education_items(lines: list[str]) -> list[EducationItem]:
    section = _section_lines(
        lines,
        ("education",),
        (
            "work experience",
            "experience",
            "employment",
            "professional experience",
            "skills",
            "core expertise",
        ),
    )
    items: list[EducationItem] = []
    for line in section:
        if _skip_resume_noise(line) or not re.search(r"\b(19\d{2}|20\d{2})\b", line):
            continue
        item = _parse_education_line(line)
        if item.school or item.degree:
            items.append(item)
    return items


def _section_lines(
    lines: list[str],
    headings: tuple[str, ...],
    stop_headings: tuple[str, ...],
) -> list[str]:
    normalized = {heading.lower().rstrip(":") for heading in headings}
    stops = {heading.lower().rstrip(":") for heading in stop_headings}
    start = -1
    for index, line in enumerate(lines):
        if line.lower().rstrip(":") in normalized:
            start = index + 1
            break
    if start < 0:
        return []
    result: list[str] = []
    for line in lines[start:]:
        normalized_line = line.lower().rstrip(":")
        if normalized_line in stops or any(normalized_line.startswith(f"{stop} ") for stop in stops):
            break
        result.append(line)
    return result


def _parse_work_header(line: str) -> WorkExperience | None:
    if line.startswith(("•", "-", "*")):
        return None
    date_parts = _find_work_date_range(line)
    if not date_parts:
        return None
    match_start, start_month, start_year, end_month, end_year = date_parts
    before = line[:match_start].strip(" -|,:")
    if len(before.split()) < 2:
        return None
    if _looks_like_description_header(before):
        return None
    title, company, location = _split_work_heading(before)
    currently_work_here = bool(end_year and end_year.lower() in {"present", "current"})
    normalized_end_month = "" if currently_work_here else (_month_number(end_month) if end_month else "12")
    normalized_end_year = "" if currently_work_here else (end_year or start_year or "")
    return WorkExperience(
        title=title,
        company=company,
        location=location,
        start_month=_month_number(start_month),
        start_year=start_year or "",
        end_month=normalized_end_month,
        end_year=normalized_end_year,
        currently_work_here=currently_work_here,
    )


def _find_work_date_range(line: str) -> tuple[int, str, str, str, str] | None:
    month_name = r"Jan|January|Feb|February|Mar|March|Apr|April|May|Jun|June|Jul|July|Aug|August|Sep|Sept|September|Oct|October|Nov|November|Dec|December"
    separator = r"\s*(?:[–—-]|\bto\b)\s*"
    patterns = [
        rf"(?P<sm>0?[1-9]|1[0-2])\s*/\s*(?P<sy>19\d{{2}}|20\d{{2}})(?:{separator}(?:(?P<em>0?[1-9]|1[0-2])\s*/\s*)?(?P<ey>19\d{{2}}|20\d{{2}}|present|current)?)?",
        rf"(?P<sm>{month_name})\s+(?P<sy>19\d{{2}}|20\d{{2}}){separator}(?:(?P<em>{month_name})\s+)?(?P<ey>19\d{{2}}|20\d{{2}}|present|current)\b",
        rf"\b(?P<sy>19\d{{2}}|20\d{{2}}){separator}(?:(?P<em>{month_name})\s+)?(?P<ey>19\d{{2}}|20\d{{2}}|present|current)\b",
        rf"\b(?P<sy>19\d{{2}}|20\d{{2}})\s*(?:[–—-]\s*)?(?P<ey>present|current)\b",
        rf"\b(?P<sy>19\d{{2}}|20\d{{2}})\b",
    ]
    for pattern in patterns:
        match = re.search(pattern, line, flags=re.IGNORECASE)
        if not match:
            continue
        groups = match.groupdict()
        return (
            match.start(),
            groups.get("sm") or "",
            groups.get("sy") or "",
            groups.get("em") or "",
            groups.get("ey") or "",
        )
    return None


def _looks_like_description_header(value: str) -> bool:
    lowered = value.strip().lower()
    return lowered.startswith(
        (
            "the research",
            "designed ",
            "developed ",
            "collaborated ",
            "awarded ",
            "selected ",
            "resulting ",
        )
    )


def _split_work_heading(value: str) -> tuple[str, str, str]:
    chunks = [_clean_work_chunk(chunk) for chunk in value.split(",")]
    chunks = [chunk for chunk in chunks if chunk]
    if len(chunks) >= 3 and _looks_like_geo_suffix(chunks[-1]):
        title, company = _split_work_title_company(", ".join(chunks[:-1]))
        return title, company, chunks[-1]
    title, company = _split_work_title_company(value)
    return title, company, ""


def _split_work_title_company(value: str) -> tuple[str, str]:
    chunks = [_clean_work_chunk(chunk) for chunk in value.split(",")]
    chunks = [chunk for chunk in chunks if chunk]
    if len(chunks) >= 3:
        if _looks_like_company(chunks[-1]) and all(
            _looks_like_work_context_chunk(chunk) for chunk in chunks[1:-1]
        ):
            return chunks[0], chunks[-1]
        return ", ".join(chunks[:-1]), chunks[-1]
    if len(chunks) == 2:
        return chunks[0], chunks[1]
    parts = re.split(r"\s{2,}|\s+\|\s+", value)
    parts = [_clean_work_chunk(part) for part in parts]
    parts = [part for part in parts if part]
    if len(parts) >= 2:
        return parts[0], parts[1]
    return _clean_work_chunk(value), ""


def _clean_work_chunk(value: str) -> str:
    cleaned = re.sub(r"\s+", " ", value).strip(" -|,:;")
    return re.sub(r"\s+\d{1,2}\s*/\s*$", "", cleaned).strip(" -|,:;")


def _looks_like_geo_suffix(value: str) -> bool:
    if _looks_like_location(value):
        return True
    return value.lower() in {
        "usa",
        "united states",
        "united states of america",
        "vietnam",
        "viet nam",
        "canada",
        "china",
        "taiwan",
        "japan",
        "south korea",
        "korea",
        "singapore",
        "malaysia",
        "india",
        "germany",
        "france",
        "united kingdom",
        "uk",
        "netherlands",
        "italy",
        "switzerland",
        "australia",
    }


def _looks_like_company(value: str) -> bool:
    return bool(
        re.search(
            r"\b(?:Corporation|Corp\.?|Incorporated|Inc\.?|LLC|Ltd\.?|Company|Co\.?)\b",
            value,
            flags=re.IGNORECASE,
        )
    )


def _looks_like_work_context_chunk(value: str) -> bool:
    lowered = value.lower()
    if _looks_like_location(value):
        return True
    if re.search(r"\b(?:division|department|group|team|unit|business|metrology)\b", lowered):
        return True
    return False


def _parse_education_line(line: str) -> EducationItem:
    year_match = re.search(r"\b(19\d{2}|20\d{2})\b", line)
    end_year = year_match.group(1) if year_match else ""
    without_year = re.sub(r"\b(19\d{2}|20\d{2})\b.*$", "", line).strip(" -|,")
    parts = [part.strip() for part in without_year.split("|") if part.strip()]
    if len(parts) >= 2:
        degree = _normalize_degree(parts[0])
        field = parts[1]
        school = parts[-1]
        if len(parts) >= 4 and _looks_like_honor(school):
            school = parts[-2]
        school = re.sub(r",\s*(?:USA|United States|Vietnam).*$", "", school).strip()
        return EducationItem(
            school=school,
            degree=degree,
            field_of_study=field,
            end_year=end_year,
        )
    return EducationItem(school=without_year, end_year=end_year)


def _normalize_degree(value: str) -> str:
    lowered = value.lower()
    if "phd" in lowered or "ph.d" in lowered or "doctor" in lowered:
        return "Doctorate"
    if lowered in {"ms", "m.s.", "msc"} or "master" in lowered:
        return "Master's Degree"
    if lowered in {"bs", "b.s.", "ba", "b.a."} or "bachelor" in lowered:
        return "Bachelor's Degree"
    if "certificate" in lowered:
        return "Certificate"
    return value


def _looks_like_honor(value: str) -> bool:
    return bool(re.search(r"\b(?:valedictorian|top ten|honors?|summa|magna|cum laude)\b", value, re.IGNORECASE))


def _add_current_work_fields(fields: dict[str, str], experience: WorkExperience) -> None:
    fields["current_job_title"] = experience.title
    fields["current_company"] = experience.company
    fields["current_job_location"] = experience.location
    fields["current_job_start_month"] = experience.start_month or "01"
    fields["current_job_start_year"] = experience.start_year


def _add_top_education_fields(fields: dict[str, str], education: EducationItem) -> None:
    fields["education_school"] = education.school
    fields["education_degree"] = education.degree
    fields["education_field"] = education.field_of_study
    fields["education_end_year"] = education.end_year


def _skip_resume_noise(line: str) -> bool:
    return bool(URL_RE.search(line) or EMAIL_RE.search(line) or re.search(r"\bPage\s+\d+\s*/\s*\d+\b", line, re.IGNORECASE))


def _append_description(existing: str, line: str) -> str:
    cleaned = re.sub(r"^[•*\-]\s*", "", line).strip()
    if not cleaned:
        return existing
    joined = f"{existing}\n{cleaned}" if existing else cleaned
    return joined[:1800]


def _month_number(value: str | None) -> str:
    if not value:
        return "01"
    stripped = value.strip().strip("/")
    if stripped.isdigit():
        month = int(stripped)
        if 1 <= month <= 12:
            return str(month).zfill(2)
    months = {
        "jan": "01",
        "january": "01",
        "feb": "02",
        "february": "02",
        "mar": "03",
        "march": "03",
        "apr": "04",
        "april": "04",
        "may": "05",
        "jun": "06",
        "june": "06",
        "jul": "07",
        "july": "07",
        "aug": "08",
        "august": "08",
        "sep": "09",
        "sept": "09",
        "september": "09",
        "oct": "10",
        "october": "10",
        "nov": "11",
        "november": "11",
        "dec": "12",
        "december": "12",
    }
    return months.get(stripped.lower(), "01")


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages).strip()


def _read_docx(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(part for part in parts if part.strip()).strip()


def _first_match(pattern: re.Pattern[str], text: str) -> str:
    match = pattern.search(text)
    return match.group(0).strip() if match else ""


def _guess_name(lines: list[str]) -> str:
    for line in lines[:6]:
        candidate = _clean_name_candidate(line)
        if not candidate:
            continue
        if _looks_like_location(candidate):
            continue
        countable = re.sub(r",?\s*(?:PhD|Ph\\.D\\.|MSc|MS|MBA|PE)$", "", candidate).strip()
        words = countable.replace(",", "").replace("(", " ").replace(")", " ").split()
        if 2 <= len(words) <= 7 and all(any(char.isalpha() for char in word) for word in words):
            return candidate
    return ""


def _clean_name_candidate(line: str) -> str:
    cleaned = URL_RE.sub(" ", line)
    cleaned = EMAIL_RE.sub(" ", cleaned)
    cleaned = PHONE_RE.sub(" ", cleaned)
    cleaned = cleaned.replace("| LinkedIn", " ")
    cleaned = cleaned.split("•", 1)[0]
    cleaned = re.sub(r"\bPage\s+\d+\s*/\s*\d+\b", " ", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" -|")
    return cleaned


def _split_name(full_name: str) -> tuple[str, str, str]:
    if not full_name:
        return "", "", ""
    without_suffix = re.sub(r",?\s*(?:PhD|Ph\\.D\\.|MSc|MS|MBA|PE)$", "", full_name).strip()
    preferred = re.search(r"\(([^)]+)\)", without_suffix)
    preferred_name = _normalize_person_name(preferred.group(1).strip()) if preferred else ""
    legal_name = re.sub(r"\([^)]*\)", " ", without_suffix)
    words = [word.strip(",") for word in legal_name.split() if word.strip(",")]
    first_name = _normalize_person_name(" ".join(words[:-1])) if len(words) > 1 else (_normalize_person_name(words[0]) if words else "")
    last_name = _normalize_person_name(words[-1]) if len(words) > 1 else ""
    return first_name, last_name, preferred_name


def _compose_legal_name(first_name: str, last_name: str) -> str:
    if not first_name and not last_name:
        return ""
    return f"{first_name} {last_name}".strip()


def _normalize_person_name(value: str) -> str:
    if not value:
        return ""
    cleaned = re.sub(r"\s+", " ", value).strip()
    if cleaned.isupper() or cleaned.islower():
        return " ".join(_normalize_name_word(word) for word in cleaned.split())
    return cleaned


def _normalize_name_word(word: str) -> str:
    parts = re.split(r"([-'])", word)
    return "".join(part.capitalize() if part not in {"-", "'"} else part for part in parts)


def _guess_section(lines: list[str], headings: tuple[str, ...]) -> str:
    normalized = {heading.lower().rstrip(":") for heading in headings}
    for index, line in enumerate(lines):
        if line.lower().rstrip(":") in normalized:
            return "\n".join(lines[index + 1 : index + 8])
    return ""


def _guess_location(lines: list[str]) -> str:
    state_pattern = re.compile(
        r"\b(?:AL|AK|AZ|AR|CA|CO|CT|DE|FL|GA|HI|IA|ID|IL|IN|KS|KY|LA|MA|MD|ME|MI|MN|MO|MS|MT|NC|ND|NE|NH|NJ|NM|NV|NY|OH|OK|OR|PA|RI|SC|SD|TN|TX|UT|VA|VT|WA|WI|WV|WY)\b"
    )
    for line in lines[:12]:
        if EMAIL_RE.search(line) or PHONE_RE.search(line):
            continue
        if _looks_like_location(line):
            return line
    return ""


def _looks_like_location(line: str) -> bool:
    state_pattern = re.compile(
        r"\b(?:AL|AK|AZ|AR|CA|CO|CT|DE|FL|GA|HI|IA|ID|IL|IN|KS|KY|LA|MA|MD|ME|MI|MN|MO|MS|MT|NC|ND|NE|NH|NJ|NM|NV|NY|OH|OK|OR|PA|RI|SC|SD|TN|TX|UT|VA|VT|WA|WI|WV|WY)\b"
    )
    return "," in line and bool(state_pattern.search(line))


def _split_city_state(location: str) -> tuple[str, str]:
    if "," not in location:
        return "", ""
    city, state = location.split(",", 1)
    return city.strip(), state.strip().split()[0]


def _guess_current_job(lines: list[str]) -> dict[str, str]:
    for line in lines:
        if re.search(r"\b(?:present|current)\b", line, flags=re.IGNORECASE):
            parts = [part.strip() for part in re.split(r"\s{2,}|\|", line) if part.strip()]
            first = parts[0] if parts else line
            first = re.sub(r"\s+\d{4}\s*[–-]\s*(?:present|current).*$", "", first, flags=re.IGNORECASE)
            title, parsed_company, location = _split_work_heading(first)
            company_match = re.search(r"\b([A-Z][A-Za-z0-9&.\- ]+?)\s+(?:Corporation|Corp\.?|Inc\.?|LLC|Company)\b", line)
            company = company_match.group(0).strip() if company_match else parsed_company
            start_match = re.search(r"\b(20\d{2}|19\d{2})\s*[–-]\s*(?:present|current)\b", line, flags=re.IGNORECASE)
            return {
                "current_job_title": title,
                "current_company": company,
                "current_job_start_month": "01",
                "current_job_start_year": start_match.group(1) if start_match else "",
                "current_job_location": location,
            }
    return {}


def _guess_top_education(lines: list[str]) -> dict[str, str]:
    for line in lines:
        if ("PhD" in line or "Ph.D" in line) and "|" in line and re.search(r"\b(20\d{2}|19\d{2})\b", line):
            parts = [part.strip() for part in line.split("|") if part.strip()]
            school = ""
            year = ""
            if len(parts) >= 3:
                school = re.sub(r",\s*USA.*$", "", parts[2]).strip()
            year_match = re.search(r"\b(20\d{2}|19\d{2})\b", line)
            if year_match:
                year = year_match.group(1)
            return {
                "education_school": school,
                "education_degree": "Doctorate",
                "education_field": parts[1] if len(parts) >= 2 else "Electrical Engineering",
                "education_end_year": year,
            }
    return {}
